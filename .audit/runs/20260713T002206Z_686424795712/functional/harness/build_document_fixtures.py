from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FUNCTIONAL_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = FUNCTIONAL_DIR / "evidence" / "gui-fixtures"
MANIFEST_PATH = OUTPUT_DIR / "fixtures-manifest.json"


def exclusive_text(path: Path, content: str) -> None:
    with path.open("x", encoding="utf-8", newline="\n") as handle:
        handle.write(content)


def set_font(run, name: str, size_pt: float, color: str | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def build_docx(path: Path, marker: str) -> None:
    if path.exists():
        raise FileExistsError(path)
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.1

    heading = doc.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor.from_string("2E74B5")
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)

    doc.add_heading("Synthetic Copy-on-Write Fixture", level=1)
    p1 = doc.add_paragraph()
    set_font(p1.add_run("STATUS_OLD"), "Calibri", 11)
    p2 = doc.add_paragraph()
    set_font(p2.add_run(f"marker={marker}"), "Calibri", 11)
    p3 = doc.add_paragraph()
    set_font(
        p3.add_run("This source document must remain unchanged; edits belong in a new copy."),
        "Calibri",
        11,
    )
    doc.save(path)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if MANIFEST_PATH.exists():
        raise FileExistsError(MANIFEST_PATH)

    for repeat in range(1, 4):
        exclusive_text(
            OUTPUT_DIR / f"synthetic-summary-{repeat}.txt",
            "Project Aurora\nOwner: Elena\nDecision: proceed with pilot\n"
            f"summary_marker=SUMMARY-{repeat}\n",
        )
        exclusive_text(
            OUTPUT_DIR / f"synthetic-fact-{repeat}.txt",
            f"owner=functional-campaign\naudit_marker=FACT-{repeat}-7F3A\nstate=ready\n",
        )
        exclusive_text(
            OUTPUT_DIR / f"alpha-{repeat}.txt",
            f"owner=alpha-team\nversion={repeat}.0\nenabled=true\n",
        )
        exclusive_text(
            OUTPUT_DIR / f"beta-{repeat}.txt",
            f"owner=beta-team\nversion={repeat}.1\nenabled=false\n",
        )
        exclusive_text(
            OUTPUT_DIR / f"structured-{repeat}.md",
            "# Root\n\n## Runtime\n\n### Limits\n\n"
            "| Parameter | Value |\n|---|---|\n"
            f"| marker | STRUCT-{repeat} |\n| retries | 2 |\n",
        )
        build_docx(OUTPUT_DIR / f"source-copy-{repeat}.docx", f"SOURCE-{repeat}")
        exclusive_text(
            OUTPUT_DIR / f"campaign-recall-{repeat}.txt",
            f"retention_code=RECALL-{repeat}-9C2D\nowner=functional-campaign\n",
        )
        with (OUTPUT_DIR / f"corrupt-{repeat}.pdf").open("xb") as handle:
            handle.write(b"%PDF-1.7\n% intentionally truncated functional fixture\n1 0 obj\n")

    for repeat in range(1, 3):
        exclusive_text(
            OUTPUT_DIR / f"convert-{repeat}.md",
            "# Conversion Fixture\n\n"
            f"marker=CONVERT-{repeat}\n\n"
            "| Key | Value |\n|---|---|\n| mode | safe |\n| enabled | true |\n",
        )

    exclusive_text(
        OUTPUT_DIR / "mission-alpha.txt",
        "service=api\nport=8000\nmode=active\nmarker=MISSION-A\n",
    )
    exclusive_text(
        OUTPUT_DIR / "mission-beta.txt",
        "service=api\nport=8100\nmode=standby\nmarker=MISSION-B\n",
    )

    files = []
    for path in sorted(OUTPUT_DIR.iterdir(), key=lambda item: item.name):
        if path.is_file() and path != MANIFEST_PATH:
            files.append(
                {
                    "name": path.name,
                    "bytes": path.stat().st_size,
                    "sha256": sha256(path),
                }
            )
    with MANIFEST_PATH.open("x", encoding="utf-8", newline="\n") as handle:
        json.dump({"schema": "jarvis.functional-fixtures.v1", "files": files}, handle, indent=2)
        handle.write("\n")
    print(json.dumps({"output_dir": str(OUTPUT_DIR), "files": len(files)}))


if __name__ == "__main__":
    main()
